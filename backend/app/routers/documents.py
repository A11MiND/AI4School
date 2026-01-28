from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from sqlalchemy.orm import Session
from ..database import get_db
from ..models.document import Document
from ..models.document_visibility import DocumentClassVisibility
from ..models.class_model import ClassModel
from ..models.student_association import StudentClass
from ..models.user import User
from ..auth.jwt import get_current_user
import io
import pypdf
import pdfplumber
import docx
import shutil
import os
import uuid
import datetime
from typing import Optional, List
from pydantic import BaseModel
from fastapi.responses import FileResponse

router = APIRouter(
    prefix="/documents",
    tags=["documents"]
)

class FolderCreate(BaseModel):
    name: str
    parent_id: Optional[int] = None

class DocumentResponse(BaseModel):
    id: int
    title: str
    file_path: Optional[str] = None
    is_folder: bool
    parent_id: Optional[int] = None
    uploaded_by: int
    created_at: Optional[datetime.datetime] = None  # Add created_at
    visible: Optional[bool] = None
    
    class Config:
        orm_mode = True

def extract_text_from_file(file_content: bytes, filename: str) -> str:
    text = ""
    if filename.endswith(".pdf"):
        # 1. Try pypdf
        try:
            pdf_reader = pypdf.PdfReader(io.BytesIO(file_content))
            for page in pdf_reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
        except Exception as e:
            print(f"PyPDF extraction failed: {e}")
            
        # 2. Key Check: If pypdf failed or returned little text, try pdfplumber
        if len(text.strip()) < 50:
            try:
                 with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                     plumber_text = ""
                     for page in pdf.pages:
                         plumber_text += (page.extract_text() or "") + "\n"
                     
                     if len(plumber_text.strip()) > len(text.strip()):
                         text = plumber_text
            except Exception as e:
                print(f"PDFPlumber extraction failed: {e}")
            
    elif filename.endswith(".docx"):
        try:
            doc = docx.Document(io.BytesIO(file_content))
            for para in doc.paragraphs:
                text += para.text + "\n"
        except Exception:
            return ""
            
    elif filename.endswith(".txt"):
        try:
            text = file_content.decode("utf-8")
        except:
             return ""
    
    return text.strip()

def normalize_extracted_text(text: str) -> str:
    if not text:
        return ""
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    normalized = normalized.replace("-\n", "")
    paragraphs = [p.strip() for p in normalized.split("\n\n") if p.strip()]
    cleaned_paragraphs = []
    for paragraph in paragraphs:
        cleaned_paragraphs.append(" ".join(paragraph.splitlines()))
    return "\n\n".join(cleaned_paragraphs).strip()

@router.post("/create_folder")
def create_folder(
    folder: FolderCreate, 
    db: Session = Depends(get_db), 
    current_user: User = Depends(get_current_user)
):
    if current_user.role != "teacher" and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only teachers can create folders")
        
    new_folder = Document(
        title=folder.name,
        is_folder=True,
        parent_id=folder.parent_id,
        uploaded_by=current_user.id
    )
    db.add(new_folder)
    db.commit()
    db.refresh(new_folder)
    return new_folder

@router.post("/upload")
async def upload_document(
    file: UploadFile = File(...), 
    parent_id: Optional[int] = Form(None),
    db: Session = Depends(get_db), 
    current_user: User = Depends(get_current_user)
):
    if current_user.role != "teacher" and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only teachers can upload documents")
    
    content_bytes = await file.read()

    # Save file to disk
    upload_dir = "uploads"
    if not os.path.exists(upload_dir):
        os.makedirs(upload_dir)
        
    unique_filename = f"{uuid.uuid4()}_{file.filename}"
    file_path = os.path.join(upload_dir, unique_filename)
    
    with open(file_path, "wb") as f:
        f.write(content_bytes)

    extracted_text = extract_text_from_file(content_bytes, file.filename)
    extracted_text = normalize_extracted_text(extracted_text)
    
    if not extracted_text:
        extracted_text = None
    
    new_doc = Document(
        title=file.filename,
        content=extracted_text,
        file_path=file_path, # Add path
        uploaded_by=current_user.id,
        parent_id=parent_id
    )
    db.add(new_doc)
    db.commit()
    db.refresh(new_doc)
    
    return {"message": "Document uploaded successfully", "id": new_doc.id}

@router.get("/", response_model=List[DocumentResponse])
def list_documents(
    parent_id: Optional[int] = None,
    uploaded_by: Optional[int] = None,
    class_id: Optional[int] = None,
    db: Session = Depends(get_db), 
    current_user: User = Depends(get_current_user)
):
    if current_user.role not in ["teacher", "admin", "student"]:
        raise HTTPException(status_code=403, detail="Authorized personnel only")
    
    query = db.query(Document)
    
    query = query.filter(Document.is_deleted == False)

    # Filter by uploader
    if current_user.role == "teacher":
        query = query.filter(Document.uploaded_by == current_user.id)
    elif uploaded_by:
        query = query.filter(Document.uploaded_by == uploaded_by)
    
    if parent_id is not None:
        query = query.filter(Document.parent_id == parent_id)
    else:
        query = query.filter(Document.parent_id == None)

    if current_user.role == "student":
        if class_id is None:
            raise HTTPException(status_code=400, detail="class_id is required")
        enrollment = db.query(StudentClass).filter(
            StudentClass.user_id == current_user.id,
            StudentClass.class_id == class_id
        ).first()
        if not enrollment:
            raise HTTPException(status_code=403, detail="Not enrolled in class")

        query = query.join(
            DocumentClassVisibility,
            DocumentClassVisibility.document_id == Document.id
        ).filter(
            DocumentClassVisibility.class_id == class_id,
            DocumentClassVisibility.visible == True
        )

    docs = query.all()
    if class_id is not None and current_user.role in ["teacher", "admin"]:
        visibility_rows = db.query(DocumentClassVisibility).filter(
            DocumentClassVisibility.class_id == class_id,
            DocumentClassVisibility.document_id.in_([d.id for d in docs])
        ).all()
        visibility_map = {row.document_id: row.visible for row in visibility_rows}
        for doc in docs:
            doc.visible = visibility_map.get(doc.id, False)

    return docs

@router.get("/{document_id}")
def get_document(document_id: int, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
     doc = db.query(Document).filter(Document.id == document_id).first()
     if not doc:
         raise HTTPException(status_code=404, detail="Document not found")
     return doc

def _collect_descendants(db: Session, root_id: int) -> list[Document]:
    collected = []
    stack = [root_id]
    while stack:
        current_id = stack.pop()
        doc = db.query(Document).filter(Document.id == current_id).first()
        if not doc:
            continue
        collected.append(doc)
        child_ids = db.query(Document.id).filter(Document.parent_id == current_id).all()
        stack.extend([cid for (cid,) in child_ids])
    return collected


@router.delete("/{document_id}")
def delete_document(
    document_id: int,
    hard: bool = False,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    if current_user.role != "teacher" and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Not authorized")
    doc = db.query(Document).filter(Document.id == document_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    targets = _collect_descendants(db, doc.id)

    if hard:
        for target in targets:
            if target.file_path and os.path.exists(target.file_path):
                try:
                    os.remove(target.file_path)
                except Exception:
                    pass
        db.query(DocumentClassVisibility).filter(
            DocumentClassVisibility.document_id.in_([t.id for t in targets])
        ).delete(synchronize_session=False)
        for target in targets:
            db.delete(target)
        db.commit()
        return {"message": "Document deleted"}

    now = datetime.datetime.now(datetime.timezone.utc)
    for target in targets:
        target.is_deleted = True
        target.deleted_at = now
    db.commit()
    return {"message": "Document deleted"}


class VisibilityUpdate(BaseModel):
    class_id: int
    visible: bool


@router.post("/{document_id}/visibility")
def set_document_visibility(
    document_id: int,
    payload: VisibilityUpdate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    if current_user.role != "teacher" and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Not authorized")

    doc = db.query(Document).filter(Document.id == document_id, Document.is_deleted == False).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    class_row = db.query(ClassModel).filter(ClassModel.id == payload.class_id).first()
    if not class_row:
        raise HTTPException(status_code=404, detail="Class not found")
    if current_user.role == "teacher" and class_row.teacher_id != current_user.id:
        raise HTTPException(status_code=403, detail="Not your class")

    visibility = db.query(DocumentClassVisibility).filter(
        DocumentClassVisibility.document_id == document_id,
        DocumentClassVisibility.class_id == payload.class_id
    ).first()
    if visibility:
        visibility.visible = payload.visible
    else:
        visibility = DocumentClassVisibility(
            document_id=document_id,
            class_id=payload.class_id,
            visible=payload.visible
        )
        db.add(visibility)

    db.commit()
    return {"message": "Visibility updated"}

@router.get("/{doc_id}/download")
def download_document(doc_id: int, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    doc = db.query(Document).filter(Document.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
        
    if current_user.role != "teacher" and current_user.role != "admin":
         if doc.uploaded_by != current_user.id:
             raise HTTPException(status_code=403, detail="Not authorized")

    if not doc.file_path or not os.path.exists(doc.file_path):
        raise HTTPException(status_code=404, detail="File not found on server")
        
    return FileResponse(doc.file_path, filename=doc.title)
